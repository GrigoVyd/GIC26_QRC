"""
Take the pandoc-generated paper_draft_v2.docx, tighten paragraph spacing,
re-render to PDF via Word COM, and merge with the Phase 1 cover page PDF
into the final submission file.

Outputs:
  paper_draft_v2.docx  (with tight spacing)
  paper_draft_v2.pdf   (body only, ideally <= 3 pages)
  Quanties__Phase2_V1.pdf  (cover + body, ready to submit)
"""

import os
import sys
import subprocess
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.join(ROOT, "..", "..")
DOCX = os.path.join(ROOT, "paper_draft_v2.docx")
BODY_PDF = os.path.join(ROOT, "paper_draft_v2.pdf")
COVER_PDF = os.path.join(PROJECT_ROOT, "Phase 1 QRC_GIC_Quanties.pdf")
FINAL_PDF = os.path.join(ROOT, "Quanties__Phase2_V1.pdf")


def tighten_docx(path: str) -> None:
    """Set paragraph spacing, margins, and force a page break before References."""
    doc = Document(path)
    for para in doc.paragraphs:
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)  # tiny gap so paragraphs are still distinguishable
        pf.line_spacing = 1.0
    from docx.shared import Inches
    from docx.enum.text import WD_BREAK
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Force a page break right before the "References" heading so the body
    # section ends cleanly on its own pages (rubric: refs don't count toward
    # the 3-page body limit, but they must be visually separable from it).
    for para in doc.paragraphs:
        if para.text.strip() == "References":
            # Insert a page-break run at the start of the References heading
            run = para.runs[0] if para.runs else para.add_run()
            # Insert a `<w:br w:type="page"/>` before the run
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run._r.insert(0, br)
            break
    doc.save(path)
    print(f"Tightened: {path}")


def render_pdf_with_word(docx_path: str, pdf_path: str) -> int:
    """Use Word COM to convert docx -> pdf and return page count."""
    ps = (
        f'$word = New-Object -ComObject Word.Application; '
        f'$word.Visible = $false; '
        f'$doc = $word.Documents.Open("{docx_path}"); '
        f'$pages = $doc.ComputeStatistics(2); '
        f'$doc.SaveAs([ref] "{pdf_path}", [ref] 17); '
        f'$doc.Close(); $word.Quit(); '
        f'Write-Output "Pages: $pages"'
    )
    out = subprocess.run(
        ["powershell", "-NoProfile", "-Command", ps],
        capture_output=True, text=True, timeout=120,
    )
    print(out.stdout.strip())
    for line in out.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":")[1].strip())
    return -1


def merge_cover_and_body(cover_pdf: str, body_pdf: str, out_pdf: str,
                          cover_pages: tuple[int, ...] = (0,)) -> None:
    """Prepend selected pages of cover_pdf to body_pdf into out_pdf."""
    from pypdf import PdfReader, PdfWriter
    writer = PdfWriter()
    cov = PdfReader(cover_pdf)
    for p in cover_pages:
        writer.add_page(cov.pages[p])
    body = PdfReader(body_pdf)
    for page in body.pages:
        writer.add_page(page)
    with open(out_pdf, "wb") as f:
        writer.write(f)
    print(f"Merged: {out_pdf}  ({len(writer.pages)} pages total)")


if __name__ == "__main__":
    tighten_docx(DOCX)
    pages = render_pdf_with_word(DOCX, BODY_PDF)
    print(f"Body PDF page count: {pages}")
    merge_cover_and_body(COVER_PDF, BODY_PDF, FINAL_PDF, cover_pages=(0,))
    # Report total pages
    from pypdf import PdfReader
    total = len(PdfReader(FINAL_PDF).pages)
    print(f"Final submission PDF: {total} pages "
          f"({total - 1} body + 1 cover, refs counted in body)")
