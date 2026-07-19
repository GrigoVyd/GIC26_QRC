"""Build the five-page GIC 2026 Phase 3 technical write-up draft.

Run this script with the Codex bundled document runtime after generating the
figures with ``build_phase3_report_figures.py``.  The output deliberately omits
the official GIC cover page because the repository currently contains a filled
Phase 2 cover, while Phase 3 requires the official Phase 3 template.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "output" / "report" / "assets"
OUT = ROOT / "output" / "report" / "Quanties_GIC2026_Phase3_Writeup_Draft.docx"

FONT = "Times New Roman"
NAVY = "17365D"
BLUE = "2F75B5"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF7E6"
WHITE = "FFFFFF"
BLACK = "000000"
RED = "9B1C1C"
GREEN = "2F6B4F"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=11, *, bold=None, italic=None, color=BLACK) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size=11, color=BLACK) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def add_text(doc, text: str, *, bold_prefix: str | None = None, after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text: str, *, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text: str, *, after=2):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(doc, text: str, level=1, *, before=None, after=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 11}[level], bold=True, color=NAVY if level < 3 else BLUE)
    return p


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=10.5, italic=True, color=MUTED)
    return p


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_callout(doc, label: str, text: str, *, fill=PALE_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    shade_paragraph(p, fill)
    r = p.add_run(f"{label} ")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int], *, font_size=10.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], "F8FAFC")
            p = cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, len(values) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def mark_last_image_alt(doc, alt: str) -> None:
    drawing = doc.inline_shapes[-1]._inline
    doc_pr = drawing.docPr
    doc_pr.set("descr", alt)
    doc_pr.set("title", alt[:80])


def add_picture(doc, filename: str, alt: str, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(str(ASSETS / filename), width=Inches(width))
    mark_last_image_alt(doc, alt)
    return p


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0

    for level, size, before, after in ((1, 16, 8, 4), (2, 13, 6, 3), (3, 11, 4, 2)):
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level < 3 else BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.42)
        style.paragraph_format.first_line_indent = Inches(-0.22)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.0

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("GIC 2026 | qBraid - MITRE - JonesTrading | Phase 3 draft")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("Team Quanties | Page ")
    set_run_font(r, size=9, color=MUTED)
    add_page_field(fp)


def page_one(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Hardware-Competitive Quantum Reservoir Computing for SPY Volatility")
    set_run_font(r, size=20, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("A native-hardware hybrid across IQM Emerald, QuEra Aquila, and cloud Ising machines")
    set_run_font(r, size=12.5, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Track A: Financial Volatility Prediction | Team Quanties | Draft 15 July 2026")
    set_run_font(r, size=10.5, bold=True, color=BLUE)

    add_heading(doc, "Abstract", level=1, before=2)
    add_text(doc, "We forecast next-day annualized 21-day realized volatility of SPY from public daily OHLCV data (2010-2024). The classical prior is GARCH(1,1); a fixed quantum reservoir learns only a regularized correction to the GARCH log-volatility forecast. This decomposition makes zero correction exactly equal to GARCH and gives the quantum component a stringent, auditable role: extract nonlinear residual structure without replacing a strong econometric model.")
    add_text(doc, "The primary real-QPU result is a 9-qubit, native 3x3 IQM Emerald reservoir evaluated on 120 chronological days at 500 shots per circuit. Its pre-specified QREM plus affine-transfer path gives RMSE 0.00831368 versus 0.00832112 for GARCH, a 0.0895% point improvement. The moving-block-bootstrap 95% interval for the RMSE difference is [-0.0000608, +0.0000262], so we claim hardware-level competitiveness, not statistically established quantum advantage. QLIKE is 0.016538 versus 0.016476 for GARCH, while the Mincer-Zarnowitz joint p-value is 0.166 (forecast unbiasedness is not rejected).")
    add_callout(doc, "Headline.", "Across executed substrates, the hybrid stays within 0.43% RMSE of matched GARCH. Real Aquila also transfers from its local AHS simulator within 0.029% RMSE, while a 12-qubit IQM expansion improves feature fidelity but worsens forecasting. Native topology and validation-locked regularization matter more than qubit count alone.")

    add_heading(doc, "1. Problem framing and data", level=1)
    add_text(doc, "The target is next-day realized volatility RV[t], computed as the rolling 21-day standard deviation of log returns multiplied by sqrt(252). At forecast time, 20 of the 21 squared-return terms are already known, which gives GARCH a structural advantage. Inputs are causal lags of returns, absolute returns, squared returns, and heterogeneous autoregressive realized-volatility summaries. All splits are chronological; model selection uses training-only expanding folds, never the held-out evaluation rows.")
    add_text(doc, "We report RMSE, volatility-specific QLIKE loss, and Mincer-Zarnowitz forecast calibration. Classical comparators include Persistence, Ridge/AR, ESN-200, LSTM, and GARCH(1,1). Two evaluation views are kept separate: a 746-day full benchmark for classical and simulator evidence, and a 400-train/120-test common window for hardware-ready and real-hardware transfer experiments.")


def page_two(doc: Document) -> None:
    add_heading(doc, "2. Hybrid QRC architecture", level=1, before=0)
    add_text(doc, "For each market state x[t], GARCH produces a log-volatility prior. A fixed reservoir maps x[t] to shot-estimated one- and two-body observables; only a Ridge readout is trained. The final forecast is log RV_hat[t] = log RV_GARCH[t] + c f_QRC(x[t]), where the Ridge penalty and correction strength c are locked on validation data. Quantum parameters are random-fixed or selected before the test window; no variational quantum optimization is used.")

    add_picture(doc, "phase3_hybrid_architecture.png", "Hybrid volatility forecasting architecture combining a GARCH prior, fixed physical quantum reservoir, shot observables, transfer correction, and Ridge residual head.", width=6.45)
    add_caption(doc, "Figure 1. The GARCH prior protects baseline performance; the fixed reservoir supplies nonlinear residual features. Hardware mitigation is label-free.")

    add_table(
        doc,
        ["Substrate", "Native encoding and dynamics", "Readout", "Purpose in study"],
        [
            ["IQM Emerald", "Rz input encoding on a selected nearest-neighbor grid; fixed CZ mixing", "45 Z/ZZ observables", "Primary real gate-QPU forecast"],
            ["QuEra Aquila", "Negative local detuning on a valid irregular atom grid; Rydberg AHS evolution", "5 Z + 10 ZZ", "Real analog-QPU transfer test"],
            ["Amplify AE / Toshiba", "Input fields h_i with fixed signed all-to-all J_ij; cloud optimization/sampling", "10 Z + 45 ZZ", "Classical signed-Ising ablation"],
            ["TFIM simulator", "Signed J_ij plus transverse-field evolution", "Z/ZZ, 3-seed ensemble", "Mechanism and advantage candidate"],
        ],
        [1500, 3300, 1500, 3060],
        font_size=9.5,
    )

    add_heading(doc, "Why a reservoir is appropriate", level=2, before=3)
    add_text(doc, "Volatility is nonlinear, heteroskedastic, and regime switching. A fixed interacting quantum system supplies a high-dimensional nonlinear projection and fading memory, while the linear head controls variance. The residual formulation is more credible than replacing GARCH: it tests whether the reservoir adds orthogonal structure after the known 21-day decomposition is supplied. Z and ZZ observables expose local response and interaction-mediated correlations.")
    add_text(doc, "Hardware transfer is deliberately conservative. IQM circuits use explicit native qubit selection and reject SWAPs; Aquila geometry satisfies the 4 micrometer per-axis separation rule and uses the device's experimental local-detuning capability. QREM and affine feature alignment use calibration information or unlabeled features only, preventing target leakage.")


def page_three(doc: Document) -> None:
    add_heading(doc, "3. Experimental design and simulator evidence", level=1, before=0)
    add_text(doc, "The full 746-day benchmark establishes task difficulty. GARCH leads the raw-task baselines (RMSE 0.007948), followed by Persistence (0.009408), ESN-200 (0.009704), and LSTM (0.010662). On the GARCH-residual target, the 10-qubit transverse-field Ising reservoir reaches RMSE 0.007844 and QLIKE 0.006445 versus 0.006863 for GARCH: improvements of 1.31% and 6.09%, respectively. The identical Ridge residual ablation has RMSE 0.008051, isolating value to the nonlinear reservoir. This is a simulator result and is not presented as executed hardware advantage.")

    add_picture(doc, "phase3_simulator_evidence.png", "Two-panel result figure showing the full classical and simulated TFIM benchmark and the improvement of four hardware-ready reservoir simulations over a linear residual ablation.", width=6.35)
    add_caption(doc, "Figure 2. Left: the simulated TFIM hybrid is the advantage candidate. Right: every hardware-ready reservoir improves on the same linear residual ablation, although none beats GARCH on this common window.")

    add_heading(doc, "Selection, noise, and ablation controls", level=2, before=3)
    add_bullet(doc, "All hyperparameters, physical layouts, Ridge penalties, and correction strengths are chosen on expanding training/validation folds. Test rows are used once for scoring.")
    add_bullet(doc, "The locked 9-qubit Rz grid candidate improves all four validation folds; its exact-statevector test gain is 0.536% in the final hardware configuration and its IQM-matched 500-shot proxy gain is 0.410%.")
    add_bullet(doc, "On the 400/120 common window, Pasqal analog, gate QRC, QuEra analog, and signed-Ising reservoirs improve RMSE over the linear residual ablation by 2.36%, 2.10%, 1.79%, and 1.52%, respectively.")
    add_bullet(doc, "The 12-qubit native-grid test is a pre-declared scaling check. It raises hardware-to-exact feature correlation from 0.770 to 0.818 but moves the forecast from 0.090% better than GARCH to 0.196% worse.")

    add_heading(doc, "Interpretation of the simulated mechanism", level=2, before=3)
    add_text(doc, "Removing the transverse field leaves the Amplify/Toshiba signed-Ising tier near GARCH but does not reproduce the full simulated gain; removing programmable signed couplings produces the neutral-atom tier with the same conclusion. The combined evidence supports a falsifiable mechanism hypothesis - useful residual features arise from the interaction of transverse-field mixing and signed couplings - while the real-QPU evidence supports competitiveness and portability rather than advantage.")


def page_four(doc: Document) -> None:
    add_heading(doc, "4. Executed hardware and cloud results", level=1, before=0)
    add_text(doc, "All reported rows below are completed external executions. Each is compared with GARCH evaluated on the identical held-out rows; therefore percentages are comparable within a row even when absolute RMSE differs across windows.")

    add_picture(doc, "phase3_hardware_evidence.png", "Two-panel hardware result figure showing RMSE gaps to matched GARCH across IQM, QuEra, Amplify, and Toshiba, plus the 9-to-12-qubit IQM scaling result.", width=6.35)
    add_caption(doc, "Figure 3. A negative RMSE gap is better than GARCH. The small 9-qubit point gain is not statistically significant; width alone does not improve the forecast.")

    add_table(
        doc,
        ["Platform", "Native workload", "Evaluation", "Budget used", "RMSE gap", "Reading"],
        [
            ["IQM Emerald 9q", "depth 32; 24 CZ; 0 SWAP", "120 x 500 shots", "20.25 Resonance cr", "-0.0895%", "Primary QPU; CI crosses zero"],
            ["IQM Emerald 12q", "depth 38; 34 CZ; 0 SWAP", "120 x 500 shots", "20.25 Resonance cr", "+0.1960%", "Scaling negative result"],
            ["QuEra Aquila 5 atoms", "native AHS; local detuning", "23 tasks x 50 shots; 20 scored", "1,840 qBraid cr", "+0.1823%", "Only 0.029% behind local AHS"],
            ["Amplify AE", "10 spins; 45 signed couplers", "120 cloud solves", "existing AE credits", "+0.0551%", "Optimizer-matched classical ablation"],
            ["Toshiba SQBM+", "10 spins; 45 signed couplers", "120 x 200 states", "existing SQBM credits", "+0.4206%", "Classical sampling ablation"],
        ],
        [1250, 2060, 1400, 1350, 900, 2400],
        font_size=9.0,
    )

    add_heading(doc, "Primary IQM result", level=2, before=2)
    add_text(doc, "Emerald job 019f5691-063b-7a11-8bbc-db5057b79259 executed 60,000 shots on physical qubits QB17-19, QB25-27, and QB33-35. QREM plus feature-only affine transfer gives feature correlation 0.770 to exact and RMSE 0.00831368. The bootstrap probability that the model is better than GARCH is 0.658; the interval crosses zero. QLIKE is slightly worse than GARCH, so the defensible conclusion is GARCH-level real-QPU performance with a small positive RMSE point estimate.")

    add_heading(doc, "Analog and cloud-Ising transfer", level=2, before=2)
    add_text(doc, "Aquila completed 23 paid tasks after two zero-cost validation failures exposed native geometry and local-detuning constraints. Three unlabeled rows calibrate feature alignment and 20 rows are scored. Hardware RMSE is 0.0137535 versus 0.0137495 for the matched local AHS hybrid and 0.0137285 for GARCH; it improves 6.06% over the Ridge residual ablation. Amplify AE and Toshiba remain within 0.43% of GARCH but are classical optimizers, not evidence of quantum advantage.")


def page_five(doc: Document) -> None:
    add_heading(doc, "5. Conclusions, impact, and reproducibility", level=1, before=0)
    add_heading(doc, "Conclusions", level=2, before=1)
    add_bullet(doc, "The strongest executed claim is cross-platform, GARCH-competitive hybrid QRC: IQM is 0.0895% better by RMSE point estimate, Aquila is 0.1823% behind GARCH and effectively matches its local AHS model, and all executed tiers remain within 0.43%.")
    add_bullet(doc, "The strongest mechanism claim remains simulated: the transverse-field signed-coupling reservoir improves RMSE and QLIKE over GARCH, whereas neutral-atom and classical signed-Ising ablations do not reproduce the full gain.")
    add_bullet(doc, "More qubits are not automatically better. Native connectivity, noise-aware training, low-depth encoding, and conservative correction strength dominate width for this data regime.")

    add_heading(doc, "Stakeholder relevance", level=2, before=3)
    add_text(doc, "Daily volatility forecasts support trading-desk hedging, market-maker inventory and spread decisions, portfolio risk limits, and scenario generation. The practical contribution is a modular residual feature engine: it can be switched off to recover GARCH exactly, deployed on multiple substrates, and monitored by its correction magnitude and calibration metrics. That is safer for financial use than an opaque end-to-end replacement.")

    add_heading(doc, "Limitations", level=2, before=3)
    add_bullet(doc, "The IQM RMSE advantage is small and not statistically significant; QLIKE does not improve. Aquila has only 20 scored days and 50 shots per task.")
    add_bullet(doc, "Amplify AE and Toshiba are useful control substrates but lack coherent transverse-field dynamics. The GARCH-beating TFIM result is simulated, not QPU-executed.")
    add_bullet(doc, "The baseline panel should ultimately include GJR-GARCH, EGARCH, and HAR-RV on the exact hardware window; current full-window baselines already include GARCH, ESN, LSTM, Ridge/AR, and Persistence.")

    add_heading(doc, "Reproducibility and submission status", level=2, before=3)
    add_table(
        doc,
        ["Requirement", "Evidence in repository", "Status"],
        [
            ["Concrete qBraid execution", "Aquila task IDs/checkpoint; qir-sv execution records", "Complete"],
            ["Re-runnable code and outputs", "Judge notebook, experiments/, compact CSV/JSON evidence", "Complete"],
            ["5-page, 11-point write-up", "This five-page body plus separate references", "Draft complete"],
            ["Official cover and Launch on qBraid", "README Launch button added; current cover is Phase 2", "Cover pending"],
        ],
        [1900, 5500, 1960],
        font_size=9.2,
    )

    add_heading(doc, "Recommended final work", level=2, before=2)
    add_number(doc, "Do not spend the remaining 580 qBraid credits. If 2,500 credits are added, authorize Aquila only after a validation-locked local analog candidate beats GARCH by at least 0.30% RMSE.")
    add_number(doc, "If that gate passes, use 31 x 50-shot tasks (estimated 2,480 credits): four engineering-smoke tasks, then resume the identical checkpoint; reserve three rows for label-free transfer and score 28.")
    add_number(doc, "Prioritize free work first: exact-window GJR/EGARCH/HAR baselines, regime/bootstrap analysis, the qBraid notebook, and a deterministic data snapshot.")
    add_number(doc, "No further paid IQM, Amplify, or Toshiba run is justified unless the rubric changes. The 9q/12q comparison and 120-row cloud runs already answer the main scaling and substrate questions.")

    add_callout(doc, "Draft completion gate.", "Insert the official Phase 3 cover page without altering its template, replace the current Phase 2 cover, add the final repository commit hash, verify the qBraid Launch link, and remove this callout before submission.", fill=PALE_GOLD)


def references(doc: Document) -> None:
    add_heading(doc, "References (excluded from the five-page limit)", level=1, before=0)
    refs = [
        "qBraid, MITRE, and JonesTrading. Global Industry Challenge 2026: Phase 3 Challenge Description. 2026.",
        "Aqora. qBraid, MITRE & JonesTrading: GIC 2026 track page. https://aqora.io/challenges/global-industry-challenge-2026/tracks/gic-2026-qBraid-MITRE-JonesTrading (accessed 15 July 2026).",
        "Bollerslev, T. Generalized autoregressive conditional heteroskedasticity. Journal of Econometrics 31, 307-327 (1986).",
        "Corsi, F. A simple approximate long-memory model of realized volatility. Journal of Financial Econometrics 7, 174-196 (2009).",
        "Patton, A. J. Volatility forecast comparison using imperfect volatility proxies. Journal of Econometrics 160, 246-256 (2011).",
        "Jaeger, H. The echo state approach to analysing and training recurrent neural networks. GMD Report 148 (2001).",
        "Fujii, K. and Nakajima, K. Harnessing disordered-ensemble quantum dynamics for machine learning. Physical Review Applied 8, 024030 (2017).",
        "Martinez-Pena, R. et al. Dynamical phase transitions in quantum reservoir computing. Physical Review Letters 127, 100502 (2021).",
        "Kornjaca, M. et al. Large-scale quantum reservoir learning with an analog quantum computer. arXiv:2407.02553 (2024).",
        "Li, Q., Mukhopadhyay, C., Bayat, A., and Habibnia, A. Quantum reservoir computing for realized volatility forecasting. arXiv:2505.13933 (2025).",
        "Ahmed, O., Tennie, F., and Magri, L. Robust quantum reservoir computers for forecasting chaotic dynamics. Proceedings of the Royal Society A 481, 20250550 (2025).",
        "qBraid. Launch on qBraid badge and repository redirect parameters. https://github.com/qBraid/community/discussions/3 (accessed 15 July 2026).",
    ]
    for ref in refs:
        add_number(doc, ref, after=4)

    add_heading(doc, "Artifact disclosure", level=2, before=8)
    add_text(doc, "Generative AI tools were used for code support, debugging, figure generation, and writing assistance. The experimental design, platform choices, executed jobs, result interpretation, and final claims are the team's own. Secrets and access tokens are excluded from the repository and must be rotated after the hardware campaign.")


def main() -> None:
    missing = [p for p in (ASSETS / "phase3_hybrid_architecture.png", ASSETS / "phase3_simulator_evidence.png", ASSETS / "phase3_hardware_evidence.png") if not p.exists()]
    if missing:
        raise FileNotFoundError(f"Missing report figures: {missing}")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    page_one(doc)
    doc.add_page_break()
    page_two(doc)
    doc.add_page_break()
    page_three(doc)
    doc.add_page_break()
    page_four(doc)
    doc.add_page_break()
    page_five(doc)
    doc.add_page_break()
    references(doc)

    core = doc.core_properties
    core.title = "Hardware-Competitive Quantum Reservoir Computing for SPY Volatility"
    core.subject = "GIC 2026 Phase 3 technical write-up draft"
    core.author = "Team Quanties"
    core.keywords = "quantum reservoir computing, volatility forecasting, IQM, QuEra, qBraid"
    core.comments = "Draft generated from saved experiment artifacts; official Phase 3 cover page not included."

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
